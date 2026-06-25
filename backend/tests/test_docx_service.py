import base64
from pathlib import Path

from docx import Document

from app.document_processing.docx import DOCXTextExtractionService


def test_docx_text_extraction_service_exports_paragraphs_and_tables(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    output_dir = tmp_path / "artifacts"

    document = Document()
    document.add_paragraph("第一段内容")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "处罚项"
    table.cell(0, 1).text = "扣分"
    table.cell(1, 0).text = "旷课"
    table.cell(1, 1).text = "20"
    document.save(docx_path)

    service = DOCXTextExtractionService()
    artifacts = service.preprocess(docx_path, output_dir)

    extracted_text = artifacts.extracted_txt_path.read_text(encoding="utf-8")
    assert "第一段内容" in extracted_text
    assert "处罚项 | 扣分" in extracted_text
    assert "旷课 | 20" in extracted_text


def test_docx_text_extraction_service_exports_embedded_image_ocr(
    monkeypatch,
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "embedded-image.docx"
    image_path = tmp_path / "sample.png"
    output_dir = tmp_path / "artifacts"

    image_path.write_bytes(
        base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aGJ0AAAAASUVORK5CYII=")
    )

    document = Document()
    paragraph = document.add_paragraph("总则说明")
    paragraph.add_run().add_picture(str(image_path))
    document.add_paragraph("附则说明")
    document.save(docx_path)

    recorded_paths: list[Path] = []

    def fake_run_ocr_on_image(path: str) -> str:
        recorded_paths.append(Path(path))
        return "图片中的规则说明"

    monkeypatch.setattr("app.document_processing.docx.service.run_ocr_on_image", fake_run_ocr_on_image)
    monkeypatch.setattr("app.document_processing.docx.service.looks_like_ocr_table", lambda _: False)

    service = DOCXTextExtractionService()
    artifacts = service.preprocess(docx_path, output_dir)

    extracted_text = artifacts.extracted_txt_path.read_text(encoding="utf-8")
    assert recorded_paths
    assert recorded_paths[0].parent == output_dir / "embedded_images"
    assert recorded_paths[0].suffix.lower() == ".png"
    assert extracted_text.index("总则说明") < extracted_text.index("图片中的规则说明") < extracted_text.index("附则说明")


def test_docx_text_extraction_service_renders_table_like_embedded_image_ocr(
    monkeypatch,
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "embedded-table.docx"
    image_path = tmp_path / "table.png"
    output_dir = tmp_path / "artifacts"

    image_path.write_bytes(
        base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aGJ0AAAAASUVORK5CYII=")
    )

    document = Document()
    document.add_paragraph().add_run().add_picture(str(image_path))
    document.save(docx_path)

    monkeypatch.setattr(
        "app.document_processing.docx.service.run_ocr_on_image",
        lambda _: "课堂纪律\n1旷课\n扣20分/次",
    )
    monkeypatch.setattr("app.document_processing.docx.service.looks_like_ocr_table", lambda _: True)
    monkeypatch.setattr(
        "app.document_processing.docx.service.render_image_ocr_table_text",
        lambda _: "[条目 1]\n内容: 旷课\n结果: 扣20分/次",
    )

    service = DOCXTextExtractionService()
    artifacts = service.preprocess(docx_path, output_dir)

    extracted_text = artifacts.extracted_txt_path.read_text(encoding="utf-8")
    assert "[条目 1]\n内容: 旷课\n结果: 扣20分/次" in extracted_text
